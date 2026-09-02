"""Requirement-document text extraction for issue intake (real file upload).

The intake write (contract v0.3 §1) keeps accepting plain
``requirement_text``; this module is how a client turns a *document* into
that text. Everything downstream continues to consume text only — the rest
of the pipeline never learns what a .docx or a .pdf is.

Supported formats are the common set: plain text / Markdown, Word (.docx),
PDF, OpenDocument (.odt) and Rich Text (.rtf). Everything else is refused
with the supported list in the message so a caller never has to guess. Legacy
.doc is deliberately absent: python-docx does not read it and translating it
would need a system binary, which a container should not depend on for a
feature whose output is text.

Parser notes:

- .docx / .odt are zip containers: python-docx and stdlib XML read the text
  nodes (paragraphs, headings, tables).
- .pdf uses pypdf's page text extraction. There is no OCR: a scanned PDF has
  no text layer and yields an empty result, which is refused rather than
  fabricated.
- .rtf is a control-word language; the stripper decodes ``\\uN`` escapes,
  drops commands and keeps run text. Lossy by design for tables and images —
  those are not requirement text anyway.

Bounds: the upload is read fully into memory (capped), and the extracted text
is capped to the intake write's own limit (``IssueIntakeCreate`` max_length)
so whatever this returns can always be submitted back with the issue.
"""

from __future__ import annotations

import io
import re
import xml.etree.ElementTree as ET
import zipfile
from dataclasses import dataclass
from pathlib import Path

from repomesh.shared.domain import DomainError

#: Largest upload we will read into memory. Documents are requirement text,
#: not media: 10 MiB covers long real PRDs while bounding memory and the
#: multipart body.
MAX_DOCUMENT_BYTES = 10 * 1024 * 1024

#: Cap on extracted text. Matches ``IssueIntakeCreate.requirement_text``
#: max_length (20 000) so the extracted text is always submittable as-is;
#: truncation is reported in the response, never hidden.
MAX_INTAKE_TEXT_CHARS = 20_000

SUPPORTED_EXTENSIONS: dict[str, str] = {
    ".txt": "text",
    ".md": "markdown",
    ".markdown": "markdown",
    ".docx": "docx",
    ".pdf": "pdf",
    ".odt": "odt",
    ".rtf": "rtf",
}

SUPPORTED_LABEL = ", ".join(sorted(SUPPORTED_EXTENSIONS))


class DocumentParseError(DomainError):
    """The file cannot be turned into requirement text."""


class UnsupportedDocumentFormat(DocumentParseError):
    """The extension is not in :data:`SUPPORTED_EXTENSIONS`."""


@dataclass(frozen=True)
class ParsedDocument:
    filename: str
    format: str
    text: str
    chars: int
    truncated: bool


def _decode_text(content: bytes) -> str:
    """Decode a text file with Chinese-friendly fallbacks.

    UTF-8 with BOM first (Word/Notepad ``另存为``), then GB18030 for legacy
    GBK/GB2312 documents, then latin-1 which never fails.
    """

    for encoding in ("utf-8-sig", "gb18030", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("latin-1", errors="replace")  # pragma: no cover


def _parse_docx(content: bytes) -> str:
    from docx import Document  # python-docx

    document = Document(io.BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(part for part in parts if part.strip())


def _parse_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 - one corrupt page must not sink the file
            continue
    return "\n\n".join(page for page in pages if page.strip())


def _parse_odt(content: bytes) -> str:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            xml_bytes = archive.read("content.xml")
    except zipfile.BadZipFile as error:
        raise DocumentParseError("odt file is not a valid zip archive") from error
    except KeyError as error:
        raise DocumentParseError("odt archive has no content.xml") from error

    root = ET.fromstring(xml_bytes)
    text_ns = "{urn:oasis:names:tc:opendocument:xmlns:text:1.0}"
    parts = [
        "".join(element.itertext()).strip()
        for element in root.iter()
        if element.tag in (f"{text_ns}p", f"{text_ns}h")
    ]
    return "\n".join(part for part in parts if part)


def _rtf_unicode(match: re.Match[str]) -> str:
    try:
        code = int(match.group(1))
    except ValueError:
        return ""
    return chr(code) if 0 <= code <= 0x10FFFF else ""


def _parse_rtf(content: bytes) -> str:
    text = _decode_text(content)
    # Literal backslash is escaped as \\ in RTF; collapse before stripping.
    text = text.replace("\\\\", "\\")
    text = text.replace("\\par", "\n").replace("\\line", "\n")
    # \\uN<fallback> — decode the code point, drop the single fallback char.
    text = re.sub(r"\\u(-?\d{1,6})[^\s]?", _rtf_unicode, text)
    # Control words: \\word, \\wordN, \\word -N (a trailing space belongs to it).
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", text)
    text = re.sub(r"\\'[0-9a-fA-F]{2}", "", text)
    text = text.replace("{", "").replace("}", "").replace("\\", "")
    return re.sub(r"[ \t]{2,}", " ", text).strip()


_PARSERS: dict[str, callable] = {
    ".txt": _decode_text,
    ".md": _decode_text,
    ".markdown": _decode_text,
    ".docx": _parse_docx,
    ".pdf": _parse_pdf,
    ".odt": _parse_odt,
    ".rtf": _parse_rtf,
}


def extract_document_text(filename: str, content: bytes) -> ParsedDocument:
    """Extract plain requirement text from a supported document.

    :raises UnsupportedDocumentFormat: the extension is not supported (415).
    :raises DocumentParseError: the file is too large, unreadable, or has no
        extractable text (413 / 422).
    """

    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedDocumentFormat(
            f"unsupported document format '{extension or Path(filename).name}'; "
            f"supported: {SUPPORTED_LABEL}"
        )
    if len(content) > MAX_DOCUMENT_BYTES:
        raise DocumentParseError(
            f"document too large ({len(content)} bytes); limit is {MAX_DOCUMENT_BYTES} bytes"
        )

    try:
        text = _PARSERS[extension](content)
    except DocumentParseError:
        raise
    except Exception as error:  # noqa: BLE001 - surface a readable refusal
        raise DocumentParseError(
            f"could not read {extension} document: {error}"
        ) from error

    text = text.strip()
    if not text:
        raise DocumentParseError(
            f"no extractable text in {Path(filename).name} — is it a scanned "
            "image PDF or an empty document?"
        )
    truncated = len(text) > MAX_INTAKE_TEXT_CHARS
    if truncated:
        text = text[:MAX_INTAKE_TEXT_CHARS]
    return ParsedDocument(
        filename=filename,
        format=SUPPORTED_EXTENSIONS[extension],
        text=text,
        chars=len(text),
        truncated=truncated,
    )
